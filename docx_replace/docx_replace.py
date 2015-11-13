#-*- coding: utf-8 -*-

from docx import Document
import sys
import copy

reload(sys)
sys.setdefaultencoding('utf-8')

formats = {
    # 0~4
    "client" : "파이온",
    "creator" : "박종수",
    "client_address" : "포항시 남구 효자동 산31번지 포항공과대학교 지곡연구동",
    "creator_address" : "포항시 남구 효자동 산31번지 포항공과대학교 기숙사 8동 105호",
    "client_representative" : "파이온",
    # 5~9
    "creator_representative" : "박종수",
    
    "total_money" : "100,000,000,000",
    "contract_date": "2015. 11. 12",
    "start_date" : "2015. 11. 13",
    "start_date_form1" : "2015년 11월 13일",
    # 10~14
    "end_date" : "2015. 12. 23",
    "end_date_form1" : "2015년 12월 13일",
    "duration_month" : "1",
    "duration_day" : "1",

    "project_name" : "게임 프로젝트 3D 그래픽",
    # 15~19
    "project_work_name" : "3D 그래픽",
    "project_range" : "3D 그래픽",

    "money_advance_percentage" : "30",
    "money_middle_percentage" : "30",
    "money_balance_percentage" : "30",
    # 20~24

    "money_advance" : "10,000",
    "money_middle" : "10,000",
    "money_balance" : "10,000",

    "middle_task" : "캐릭터 3종",

    "as_duration_in_month" : "10",

    # 25~29
    "contract_cancellation_reward_client2creator" : "10",
    "contract_cancellation_reward_creator2client" : "10",
}

def change_paragraph(paragraph, filter_format):
    for format, content in filter_format.iteritems():
        search_word = "{"+format+"}"
        needCheck = True if search_word in paragraph.text else False
        for run in paragraph.runs:
            if search_word in run.text:
                replaced = run.text.replace(search_word, content)
                run = run.clear()
                run.add_text(unicode(replaced))
                needCheck = False
        if needCheck:
            #print format, " is not checked in ", paragraph.text
            start_idx = 0
            end_idx = 0
            complete_word = ""
            i = 0
            broken_keys = []
            for run in paragraph.runs:
                if run.text == "{":
                    start_idx = i
                    complete_word = run.text
                else:
                    end_idx = i
                    complete_word = complete_word + run.text
                    if run.text == "}":
                        replaced = complete_word.replace(search_word, content)
                        broken_keys.append((start_idx, end_idx, replaced))
                #print ":::", run.text
                i = i+1

            for (start_idx, end_idx, replaced) in broken_keys:
                for run_idx, run in enumerate(paragraph.runs):
                    if start_idx <= run_idx and run_idx <= end_idx:
                        run = run.clear()
                    if run_idx == end_idx:
                        #print "))", replaced
                        run.add_text(unicode(replaced))

def replace(infile, outfile, filter_format):
    if filter_format == None:
        filter_format = formats

    document = Document(infile)

    for paragraph in document.paragraphs:
        change_paragraph(paragraph, filter_format)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    change_paragraph(paragraph, filter_format)
    
    document.save(outfile)
